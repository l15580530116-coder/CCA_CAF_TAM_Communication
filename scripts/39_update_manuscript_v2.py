#!/usr/bin/env python
"""Step 39: Update manuscript legends and Results text for expanded A-F panels."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import re, copy

BASE = Path(r"e:\CCA")
SRC_DOCX = BASE / "submission_CSBJ_UPLOAD_READY/Manuscript_CSBJ_final_clean.docx"
DST_DOCX = BASE / "submission_CSBJ_UPLOAD_READY/Manuscript_CSBJ_final_clean_v2.docx"

doc = Document(str(SRC_DOCX))

def set_run(run, size=11, bold=False, italic=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

# ═══════════════════════════════════════════════════════════════
# 1. REPLACE FIGURE LEGENDS
# ═══════════════════════════════════════════════════════════════
# Search through all paragraphs for Figure legends and replace them

new_legends = {
    "Figure 5.": (
        "Figure 5. CellChat-inferred CAF-TAM-Epithelial communication network. "
        "(A) Simplified directed network summarizing inferred communication among CAF_Fibroblast, "
        "Macrophage_TAM, and Epithelial_like cells. "
        "(B) Incoming and outgoing signaling patterns across tumor cell types. "
        "(C) Bubble plot of inferred CAF-to-TAM ligand-receptor pairs. "
        "(D) Bubble plot of inferred TAM-to-CAF ligand-receptor pairs. "
        "(E) Top inferred signaling pathways ranked by ligand-receptor pair count. "
        "(F) Source-target communication strength summary. "
        "CellChat results represent computational inference based on ligand-receptor co-expression "
        "and do not establish functional signaling."
    ),
    "Figure 6.": (
        "Figure 6. Integrated hub gene prioritization and exploratory clinical relevance. "
        "(A) Multi-layer evidence heatmap for candidate hub genes. "
        "(B) Ranking of top hub genes based on integrated evidence scores. "
        "(C) Prioritized ligand-receptor communication axes. "
        "(D) Exploratory association between CAF-TAM axis scores and overall survival in GSE107943. "
        "(E) External expression-level validation of hub genes in GSE26566. "
        "(F) Hub gene-score correlation summary. "
        "Survival analyses are exploratory and require validation in larger independent cohorts."
    ),
    "Figure 7.": (
        "Figure 7. Candidate target prioritization and in silico docking. "
        "(A) Druggability classification of prioritized targets. "
        "(B) Candidate drug-target relationship network. "
        "(C) Candidate drug count per target. "
        "(D) Predicted docking affinities for selected target-ligand pairs. "
        "(E) Docking neighbor residue summary. "
        "(F) Proposed model summarizing a CAF-TAM communication-associated aggressive "
        "microenvironment in cholangiocarcinoma and candidate targets for future validation. "
        "Docking scores are computational predictions and do not represent experimental binding affinities."
    ),
}

legend_replacements = 0
for para in doc.paragraphs:
    text = para.text.strip()
    for prefix, new_text in new_legends.items():
        if text.startswith(prefix):
            # Clear all runs and set new text
            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = new_text
                set_run(para.runs[0], size=11)
                # Remove extra runs
                for run in para.runs[1:]:
                    run.text = ""
            else:
                run = para.add_run(new_text)
                set_run(run, size=11)
            legend_replacements += 1
            print(f"  Replaced: {prefix}")

print(f"Legends replaced: {legend_replacements}/3")

# ═══════════════════════════════════════════════════════════════
# 2. ADD RESULTS TEXT REFERENCES
# ═══════════════════════════════════════════════════════════════
# Find paragraphs in Results section and insert new sentences

results_insertions = {
    # Results 4.5: After "TAM-to-CAF ligand-receptor pairs" sentence
    "4.5_add_EF": {
        "search": "TAM-to-CAF ligand-receptor pairs",
        "add": " The top inferred signaling pathways and source-target communication strengths are summarized in Figure 5E and Figure 5F."
    },
    # Results 4.7: After GSE26566 sentence
    "4.7_add_EF": {
        "search": "precluding survival validation in this third cohort",
        "add": " External expression-level validation of hub genes in GSE26566 and hub gene-score correlation patterns are shown in Figure 6E and Figure 6F."
    },
    # Results 4.8: After docking paragraph
    "4.8_add_BCEF": {
        "search": "co-crystal ligand centroids (Figure 7B)",
        "add": " Candidate drug-target relationships are summarized in Figure 7B, with candidate drug counts per target shown in Figure 7C. Docking neighbor residue analysis identified 43 residues within 4 Angstrom of the best docking poses across the four target-ligand pairs (Figure 7E; Supplementary Table S17). A proposed CAF-TAM-Epithelial model integrating these findings is shown in Figure 7F."
    },
}

text_insertions = 0
for para in doc.paragraphs:
    text = para.text
    for key, info in results_insertions.items():
        search = info["search"]
        add = info["add"]
        if search in text and add not in text:
            # Find the position and insert
            idx = text.find(search) + len(search)
            # Check if it ends with period - insert before period
            remaining = text[idx:]
            # Insert the new text
            new_text = text[:idx] + "." + add + remaining
            # Update runs
            if para.runs:
                # Simple approach: append to last run
                last_run = para.runs[-1]
                if add not in last_run.text:
                    last_run.text = last_run.text + ". " + add
                    text_insertions += 1
                    print(f"  Added: {key}")
            break

print(f"Text insertions: {text_insertions}/3")

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
doc.save(str(DST_DOCX))
print(f"\nSaved: {DST_DOCX.name}")
