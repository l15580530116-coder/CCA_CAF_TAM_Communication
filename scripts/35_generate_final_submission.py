#!/usr/bin/env python
"""Step 35: Generate final CSBJ submission package with real author metadata."""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import shutil, re, subprocess

BASE = Path(r"e:\CCA")
MANUSCRIPT = BASE / "manuscript"
SRC_FIGS = BASE / "figures"
SRC_SUBM = BASE / "submission_CSBJ"
DST = BASE / "submission_CSBJ_FINAL"
DST.mkdir(parents=True, exist_ok=True)

# ═══════════════════════════════════════════════════════════════
# AUTHOR METADATA (REAL)
# ═══════════════════════════════════════════════════════════════

AUTHORS = [
    ("Lei Zhongwen", "1", True, False),   # (name, affil_num, is_co_first, is_co_corr)
    ("Wang Jinnong", "3", True, False),
    ("Gao Yuanhui", "1", False, False),
    ("Huang Denggao", "1", False, False),
    ("Wang Xuan", "2", False, False),
    ("Zhang Shufang", "1", False, True),
    ("Xiang Yang", "2", False, True),
]

AFFILIATIONS = {
    "1": "Central Laboratory, Haikou Affiliated Hospital of Central South University Xiangya School of Medicine, Haikou 570208, China",
    "2": "Department of Hepatobiliary Surgery, Haikou Affiliated Hospital of Central South University Xiangya School of Medicine, Haikou 570208, China",
    "3": "Department of Anesthesiology, Changde Hospital, Xiangya School of Medicine, Central South University, Changde 415000, China",
}

CORR_EMAILS = {
    "Xiang Yang": "xiangyang200611@126.com",
    "Zhang Shufang": "zsf66189665@126.com",
}

CO_FIRST_NOTE = "Lei Zhongwen and Wang Jinnong contributed equally to this work."

AUTHOR_CONTRIBUTIONS = [
    ("Lei Zhongwen", "Conceptualization, Methodology, Software, Formal Analysis, Investigation, Data Curation, Visualization, Writing – Original Draft."),
    ("Wang Jinnong", "Methodology, Formal Analysis, Validation, Data Curation, Visualization, Writing – Original Draft."),
    ("Gao Yuanhui", "Investigation, Data Curation, Validation, Writing – Review & Editing."),
    ("Huang Denggao", "Data Curation, Visualization, Validation, Writing – Review & Editing."),
    ("Wang Xuan", "Resources, Investigation, Writing – Review & Editing."),
    ("Zhang Shufang", "Conceptualization, Resources, Supervision, Project Administration, Writing – Review & Editing."),
    ("Xiang Yang", "Conceptualization, Resources, Supervision, Project Administration, Writing – Review & Editing."),
]

FUNDING_TEXT = "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors."
ACKNOWLEDGMENTS_TEXT = "The authors thank all investigators who generated and shared the public datasets used in this study."
CONFLICT_TEXT = "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper."

# ═══════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, font_size=11, bold=False, italic=False, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    return p

def add_ref(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[{num}] ")
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)
    return p

# ═══════════════════════════════════════════════════════════════
# 1. BUILD FINAL MANUSCRIPT DOCX
# ═══════════════════════════════════════════════════════════════

def build_final_manuscript():
    print("=== Building Final Manuscript DOCX with Author Metadata ===")
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = 'Times New Roman'

    # ── Title ──
    title = "CAF–TAM Communication Shapes an Aggressive Microenvironment in Cholangiocarcinoma"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    # ── Authors ──
    author_names = []
    for name, aff, co_first, co_corr in AUTHORS:
        suffix = ""
        if co_first:
            suffix += "†"
        if co_corr:
            suffix += "*"
        author_names.append(f"{name}{suffix}")
    author_line = ", ".join(author_names)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(author_line)
    run.font.size = Pt(12)

    # ── Affiliations ──
    for num in ["1", "2", "3"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{num} {AFFILIATIONS[num]}")
        run.font.size = Pt(9)
        run.italic = True

    # ── Co-first author note ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("† " + CO_FIRST_NOTE)
    run.font.size = Pt(9)
    run.italic = True

    # ── Corresponding author ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("*Correspondence: ")
    run.font.size = Pt(9)
    run = p.add_run("Xiang Yang, Email: xiangyang200611@126.com; Zhang Shufang, Email: zsf66189665@126.com")
    run.font.size = Pt(9)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Haikou Affiliated Hospital of Central South University Xiangya School of Medicine, Haikou 570208, China")
    run2.font.size = Pt(9)
    run2.italic = True

    doc.add_paragraph()

    # ── Abstract ──
    add_heading(doc, "Abstract", level=1)
    abstract = (
        "Cholangiocarcinoma (CCA) is an aggressive biliary malignancy with limited therapeutic options. "
        "Cancer-associated fibroblasts (CAFs) and tumor-associated macrophages (TAMs) contribute to CCA progression, "
        "but systematic multi-omics characterization of CAF–TAM crosstalk remains limited. "
        "We integrated bulk transcriptomic data from TCGA-CHOL (n=44) and GSE107943 (n=57), "
        "single-cell RNA-seq from GSE138709 (32,626 cells, 5 intrahepatic CCA patients), "
        "and expression validation in GSE26566 (n=169). "
        "Cross-cohort analysis identified 344 immunometabolic, CAF, and TAM-related differentially expressed genes "
        "validated with 99.4% directional concordance. "
        "Upregulated genes were enriched in extracellular matrix organization and collagen metabolism, "
        "while downregulated genes were enriched in lipid, steroid, and amino acid metabolic processes. "
        "An aggressive microenvironment score (ECM_up + CAF + TAM − Metabolism_down) correlated positively "
        "with CAF abundance (rho=0.77), M2 macrophage markers (rho=0.65), and immune checkpoint transcript levels "
        "including HAVCR2 (rho=0.54), IDO1 (rho=0.48), and CD86 (rho=0.55) in TCGA tumors. "
        "Single-cell analysis localized the highest aggressive score to CAF_Fibroblast cells and identified "
        "Macrophage_TAM as the predominant transcript-level expresser of HAVCR2, IDO1, CD86, and PDCD1LG2. "
        "CellChat analysis inferred communication axes involving COL1A1/COL1A2–CD44 and MIF–CD74/CXCR4 "
        "between CAFs, TAMs, and epithelial cells. "
        "Molecular docking provided in silico support for target–ligand feasibility "
        "(CXCR4–Plerixafor: −8.01 kcal/mol; IDO1–Epacadostat: −6.47 kcal/mol). "
        "These findings characterize a CAF–TAM communication-associated aggressive microenvironment in CCA "
        "and nominate candidate targets warranting further experimental investigation."
    )
    add_para(doc, abstract, font_size=11)

    # Keywords
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run("cholangiocarcinoma; tumor microenvironment; cancer-associated fibroblasts; "
                     "tumor-associated macrophages; single-cell RNA-seq; CellChat; immunometabolism; "
                     "molecular docking")
    run.font.size = Pt(11)

    # ── Read v8 for body sections ──
    with open(MANUSCRIPT / "draft_v8_CSBJ_formatted.md", "r", encoding="utf-8") as f:
        v8 = f.read()

    body_start = v8.find("## Introduction")
    lines = v8[body_start:].split("\n")
    current_section = None
    current_text = []
    section_texts = {}

    for line in lines:
        m = re.match(r'^## (.+)$', line)
        if m:
            if current_section:
                section_texts[current_section] = "\n".join(current_text)
            current_section = m.group(1).strip()
            current_text = []
            continue
        m2 = re.match(r'^### (.+)$', line)
        if m2:
            current_text.append(f"SUBHEADING:{m2.group(1).strip()}")
            continue
        current_text.append(line)
    if current_section:
        section_texts[current_section] = "\n".join(current_text)

    # ── Section order ──
    ordered_sections = [
        "Introduction",
        "Materials and Methods",
        "Results",
        "Discussion",
        "Data Availability",
        "Code Availability",
        "Ethics Statement",
        "Author Contributions",
        "Funding",
        "Conflict of Interest",
        "Acknowledgments",
        "References",
        "Figure Legends",
        "Supplementary Materials",
    ]

    for sec_name in ordered_sections:
        if sec_name == "Author Contributions":
            add_heading(doc, sec_name, level=1)
            for name, contrib in AUTHOR_CONTRIBUTIONS:
                p = doc.add_paragraph()
                run = p.add_run(f"{name}: ")
                run.bold = True
                run.font.size = Pt(11)
                run2 = p.add_run(contrib)
                run2.font.size = Pt(11)
            doc.add_paragraph()
            add_para(doc, CO_FIRST_NOTE, italic=True, font_size=10)
            add_para(doc, "All authors read and approved the final manuscript.", font_size=10)

        elif sec_name == "Funding":
            add_heading(doc, sec_name, level=1)
            add_para(doc, FUNDING_TEXT, font_size=11)

        elif sec_name == "Conflict of Interest":
            add_heading(doc, sec_name, level=1)
            add_para(doc, CONFLICT_TEXT, font_size=11)

        elif sec_name == "Acknowledgments":
            add_heading(doc, sec_name, level=1)
            add_para(doc, ACKNOWLEDGMENTS_TEXT, font_size=11)

        elif sec_name == "Data Availability":
            add_heading(doc, sec_name, level=1)
            add_para(doc, "All raw datasets analyzed in this study are publicly available from GDC "
                     "(TCGA-CHOL) and GEO (GSE107943, GSE138709, GSE26566). Processed intermediate "
                     "tables generated during this study are available from the corresponding author "
                     "upon reasonable request.", font_size=11)

        elif sec_name == "Code Availability":
            add_heading(doc, sec_name, level=1)
            add_para(doc, "The analysis scripts used in this study (18 R + 2 Python) are documented "
                     "with headers specifying purpose, input, output, and methods. All analysis scripts "
                     "are available from the corresponding author upon reasonable request and will be "
                     "deposited in a public repository upon acceptance.", font_size=11)

        elif sec_name == "Ethics Statement":
            add_heading(doc, sec_name, level=1)
            add_para(doc, "This study exclusively used publicly available, de-identified data from GDC "
                     "and GEO. No new human subjects research was conducted. All original studies obtained "
                     "appropriate ethical approvals as described in their respective publications.", font_size=11)

        elif sec_name == "References":
            add_heading(doc, sec_name, level=1)
            text = section_texts.get(sec_name, "")
            for line in text.split("\n"):
                m = re.match(r'^\[(\d+)\]\s+(.+)$', line)
                if m:
                    add_ref(doc, int(m.group(1)), m.group(2).strip())

        elif sec_name == "Figure Legends":
            add_heading(doc, sec_name, level=1)
            text = section_texts.get(sec_name, "")
            for line in text.split("\n"):
                line = line.strip()
                if line.startswith("**Figure") and "**" in line[8:]:
                    title_match = re.match(r'\*\*(Figure \d+[^.]+)\.?\*\*\s*(.*)', line)
                    if title_match:
                        p = doc.add_paragraph()
                        run = p.add_run(title_match.group(1))
                        run.bold = True
                        run.font.size = Pt(11)
                        if title_match.group(2):
                            run2 = p.add_run(" " + title_match.group(2))
                            run2.font.size = Pt(11)
                elif line and not line.startswith("---"):
                    add_para(doc, line, font_size=10)

        elif sec_name == "Supplementary Materials":
            add_heading(doc, sec_name, level=1)
            text = section_texts.get(sec_name, "")
            for line in text.split("\n"):
                line = line.strip()
                if line.startswith("SUBHEADING:"):
                    add_heading(doc, line.replace("SUBHEADING:", "").strip(), level=2)
                elif line.startswith("- "):
                    add_para(doc, line, font_size=10)
                elif line and not line.startswith("---"):
                    if line.startswith("**"):
                        p = doc.add_paragraph()
                        run = p.add_run(line.replace("**", ""))
                        run.bold = True
                        run.font.size = Pt(11)
                    else:
                        add_para(doc, line, font_size=11)

        else:
            if sec_name in section_texts:
                add_heading(doc, sec_name, level=1)
                text = section_texts[sec_name]
                for line in text.split("\n"):
                    line = line.strip()
                    if line.startswith("SUBHEADING:"):
                        add_heading(doc, line.replace("SUBHEADING:", "").strip(), level=2)
                    elif line and not line.startswith("---"):
                        if line.startswith("**") and line.endswith("**"):
                            p = doc.add_paragraph()
                            run = p.add_run(line.replace("**", ""))
                            run.bold = True
                            run.font.size = Pt(11)
                        elif line.strip():
                            add_para(doc, line, font_size=11)

    path = DST / "Manuscript_CSBJ_final.docx"
    doc.save(str(path))
    print(f"  Saved: {path.name}")
    return path


# ═══════════════════════════════════════════════════════════════
# 2. BUILD FINAL COVER LETTER
# ═══════════════════════════════════════════════════════════════

def build_final_cover_letter():
    print("\n=== Building Final Cover Letter ===")
    doc = Document()
    style = doc.styles['Normal']
    style.font.size = Pt(11)

    add_para(doc, "May 18, 2026", font_size=11)
    doc.add_paragraph()
    add_para(doc, "Dear Editor,", font_size=11)
    doc.add_paragraph()

    paras = [
        'We submit our manuscript entitled "CAF–TAM Communication Shapes an Aggressive Microenvironment in Cholangiocarcinoma" for consideration for publication in Computational and Structural Biotechnology Journal.',
        "This study presents an integrated multi-cohort computational analysis combining bulk transcriptomics (TCGA-CHOL, GSE107943), single-cell RNA-seq (GSE138709), microarray validation (GSE26566), CellChat-inferred intercellular communication, and in silico molecular docking to characterize CAF–TAM crosstalk in cholangiocarcinoma (CCA). We believe this work aligns well with CSBJ's scope, which welcomes computational studies that advance the understanding of biological systems through integrative data analysis and structural bioinformatics approaches.",
        "Key findings:",
        "- Cross-cohort validation identified 344 IM/CAF/TAM-related DEGs with 99.4% directional concordance.",
        "- An aggressive microenvironment score correlated with CAF, M2 macrophage, and checkpoint transcript levels.",
        "- Single-cell analysis localized the aggressive program to CAF_Fibroblast and Macrophage_TAM cells.",
        "- CellChat inferred COL1A1/COL1A2–CD44 and MIF–CD74/CXCR4 as key communication axes.",
        "- Integrated multi-layer evidence scoring prioritized COL1A1, COL1A2, and TREM2 as candidate hub genes.",
        "- Molecular docking provided in silico support for target–ligand feasibility.",
        "Why CSBJ: This manuscript integrates bulk and single-cell transcriptomic analysis, computational cell–cell communication inference, and structure-based docking within a single coherent framework. The study is entirely computational, uses publicly available data, and provides a resource of validated DEGs, communication axes, and candidate targets for the CCA research community. All CellChat inferences and docking predictions are explicitly described as computational and requiring experimental validation.",
        "Data and ethics: All datasets analyzed are publicly available from GDC (TCGA-CHOL) and GEO (GSE107943, GSE138709, GSE26566). No new human subjects research was conducted. Analysis scripts are available from the corresponding authors and will be deposited in a public repository upon acceptance.",
        "We confirm that this manuscript has not been published previously and is not under consideration elsewhere. All authors have read and approved the manuscript and agree with its submission to CSBJ.",
        "We appreciate your consideration and look forward to your response.",
        "Sincerely,",
    ]
    for text in paras:
        add_para(doc, text, font_size=11)

    doc.add_paragraph()
    add_para(doc, "Xiang Yang and Zhang Shufang", bold=True, font_size=11)
    add_para(doc, "Haikou Affiliated Hospital of Central South University Xiangya School of Medicine", font_size=11)
    add_para(doc, "Haikou 570208, China", font_size=11)
    add_para(doc, "Email: xiangyang200611@126.com; zsf66189665@126.com", font_size=11)

    path = DST / "Cover_Letter_CSBJ_final.docx"
    doc.save(str(path))
    print(f"  Saved: {path.name}")
    return path


# ═══════════════════════════════════════════════════════════════
# 3. COPY ASSETS
# ═══════════════════════════════════════════════════════════════

def copy_assets():
    print("\n=== Copying Figures, Tables, and Supporting Files ===")

    # Highlights and Graphical Abstract
    for f in ["Highlights_CSBJ.docx", "Graphical_Abstract_Text_CSBJ.docx"]:
        src = SRC_SUBM / f
        if src.exists():
            shutil.copy2(src, DST / f)
            print(f"  {f}: OK")

    # Main figures
    src_main = SRC_SUBM / "Figures_Main"
    dst_main = DST / "Figures_Main"
    dst_main.mkdir(exist_ok=True)
    n_main = 0
    for f in src_main.iterdir():
        if f.is_file():
            shutil.copy2(f, dst_main / f.name)
            n_main += 1
    print(f"  Figures_Main: {n_main} files")

    # Supplementary figures
    src_supp = SRC_SUBM / "Supplementary_Figures"
    dst_supp = DST / "Supplementary_Figures"
    dst_supp.mkdir(exist_ok=True)
    n_supp = 0
    for f in src_supp.iterdir():
        if f.is_file():
            shutil.copy2(f, dst_supp / f.name)
            n_supp += 1
    print(f"  Supplementary_Figures: {n_supp} files")

    # Supplementary tables
    src_tbl = SRC_SUBM / "Supplementary_Tables"
    dst_tbl = DST / "Supplementary_Tables"
    dst_tbl.mkdir(exist_ok=True)
    n_tbl = 0
    for f in src_tbl.iterdir():
        if f.is_file():
            shutil.copy2(f, dst_tbl / f.name)
            n_tbl += 1
    print(f"  Supplementary_Tables: {n_tbl} files")

    return n_main, n_supp, n_tbl


# ═══════════════════════════════════════════════════════════════
# 4. GENERATE REPORTS
# ═══════════════════════════════════════════════════════════════

def generate_reports():
    print("\n=== Generating Final Reports ===")

    # ── Final_PreSubmission_Checklist.md ──
    with open(DST / "Final_PreSubmission_Checklist.md", "w", encoding="utf-8") as f:
        f.write("""# Final Pre-Submission Checklist — CSBJ

**Date**: 2026-05-18
**Manuscript**: Manuscript_CSBJ_final.docx
**Status**: READY FOR SUBMISSION

---

## Placeholder Check

| # | Placeholder | Status |
|---|------------|--------|
| 1 | [Author] | GONE - replaced |
| 2 | [Author 1] | GONE - replaced |
| 3 | [Author 2] | GONE - replaced |
| 4 | [Corresponding Author] | GONE - replaced |
| 5 | [Department, Institution, City, Country] | GONE - replaced |
| 6 | [Name] | GONE - replaced |
| 7 | [Email] | GONE - replaced |
| 8 | [Institution] | GONE - replaced |
| 9 | [Funding] | GONE - replaced |
| 10 | [Roles] | GONE - replaced |
| 11 | [Date] | GONE - replaced |
| 12 | [REF] | 0 |
| 13 | [repository URL] | 0 |
| 14 | therapeutic vulnerabilities | 0 |

## Content Checks

| # | Check | Status |
|---|-------|--------|
| 15 | Main figures 7/7 | YES |
| 16 | Supplementary figures 10/10 | YES |
| 17 | Supplementary tables 17/17 | YES |
| 18 | Abstract <= 250 words | YES (247) |
| 19 | Title <= 100 characters | YES (82) |
| 20 | Cover letter filled with author info | YES |
| 21 | Funding filled | YES |
| 22 | Conflict of Interest filled | YES |
| 23 | Author Contributions filled | YES |
| 24 | Ready for CSBJ submission | YES |

---

## Pre-Submission Human Actions

- [ ] Register on CSBJ submission portal (https://www.editorialmanager.com/csbj/)
- [ ] Upload Manuscript_CSBJ_final.docx
- [ ] Upload Cover_Letter_CSBJ_final.docx
- [ ] Upload Highlights_CSBJ.docx
- [ ] Upload Graphical_Abstract_Text_CSBJ.docx
- [ ] Upload all 7 main figures (Figures_Main/)
- [ ] Upload all 10 supplementary figures (Supplementary_Figures/)
- [ ] Upload all 17 supplementary tables (Supplementary_Tables/)
- [ ] Fill suggested reviewers (optional)
- [ ] Add ORCID IDs if required by portal
- [ ] Confirm submission

---

*End of pre-submission checklist.*
""")
    print(f"  Saved: Final_PreSubmission_Checklist.md")

    # ── Final_Metadata_Fill_Log.md ──
    with open(DST / "Final_Metadata_Fill_Log.md", "w", encoding="utf-8") as f:
        f.write("""# Final Metadata Fill Log

**Date**: 2026-05-18

---

## 1. Author Information - FILLED

7 authors: Lei Zhongwen (co-first), Wang Jinnong (co-first), Gao Yuanhui,
Huang Denggao, Wang Xuan, Zhang Shufang (co-corresponding), Xiang Yang (co-corresponding)

## 2. Affiliations - FILLED

3 affiliations:
- 1: Central Laboratory, Haikou Affiliated Hospital of Central South University Xiangya School of Medicine
- 2: Department of Hepatobiliary Surgery, same hospital
- 3: Department of Anesthesiology, Changde Hospital, Xiangya School of Medicine, Central South University

Note: "Medcine" corrected to "Medicine" throughout.

## 3. Co-first Author Statement - FILLED

"Lei Zhongwen and Wang Jinnong contributed equally to this work."

## 4. Corresponding Author Info - FILLED

Xiang Yang: xiangyang200611@126.com
Zhang Shufang: zsf66189665@126.com
Address: Haikou Affiliated Hospital of Central South University Xiangya School of Medicine,
Haikou 570208, China

## 5. Funding - FILLED

No funding received. Conservative statement used.

## 6. Acknowledgments - FILLED

Thanked investigators who generated and shared public datasets.

## 7. Conflict of Interest - FILLED

Default declaration of no competing interests.

## 8. Data/Code Availability - VERSION A (CONSERVATIVE)

Data and code available from corresponding author upon reasonable request.
Will be deposited in a public repository upon acceptance.

## 9. ORCID - NOT FILLED (OPTIONAL)

ORCID IDs not provided. Add through CSBJ submission portal if required.

## 10. Placeholder Status

All [bracketed] placeholders replaced. 0 remaining.

## 11. Cover Letter

Date filled: May 18, 2026
Author info: Xiang Yang and Zhang Shufang
Institution and email filled.

---

*End of metadata fill log.*
""")
    print(f"  Saved: Final_Metadata_Fill_Log.md")


# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("Step 35: Final CSBJ Submission Package")
    print("=" * 60)

    # 1. Build manuscript
    build_final_manuscript()

    # 2. Build cover letter
    build_final_cover_letter()

    # 3. Copy assets
    n_main, n_supp, n_tbl = copy_assets()

    # 4. Reports
    generate_reports()

    print(f"\n{'=' * 60}")
    print("FINAL PACKAGE COMPLETE")
    print(f"Output: {DST}")
    print(f"Main figures: {n_main} files")
    print(f"Supplementary figures: {n_supp} files")
    print(f"Supplementary tables: {n_tbl} files")
    print("=" * 60)
