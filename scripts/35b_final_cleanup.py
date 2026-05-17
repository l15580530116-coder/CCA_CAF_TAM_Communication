#!/usr/bin/env python
"""Step 35b: Fix affiliation numbers, Figure Legends, and S17 reference."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
import re, shutil, copy

BASE = Path(r"e:\CCA")
MANUSCRIPT = BASE / "manuscript"
SRC = BASE / "submission_CSBJ_FINAL"
DST = BASE / "submission_CSBJ_FINAL"

def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)

def add_para(doc, text, font_size=11, bold=False, italic=False, alignment=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if alignment is not None:
        p.alignment = alignment
    return p

def add_ref(doc, num, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"[{num}] ")
    run.font.size = Pt(10)
    run2 = p.add_run(text)
    run2.font.size = Pt(10)

def add_superscript_run(paragraph, text, font_size=9, bold=False):
    """Add a superscript run to a paragraph."""
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.superscript = True
    return run

def add_run(paragraph, text, font_size, bold=False, italic=False):
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return run

# ═══════════════════════════════════════════════════════════════
# BUILD CLEAN MANUSCRIPT
# ═══════════════════════════════════════════════════════════════

print("=== Building Clean Final Manuscript ===")
doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Times New Roman'

# ── Title ──
title = "CAF–TAM Communication Shapes an Aggressive Microenvironment in Cholangiocarcinoma"
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(title)
run.bold = True
run.font.size = Pt(14)

# ── Authors with superscript affiliation numbers ──
authors_data = [
    ("Lei Zhongwen", "1", True, False),   # (name, aff_num, co_first, co_corr)
    ("Wang Jinnong", "3", True, False),
    ("Gao Yuanhui", "1", False, False),
    ("Huang Denggao", "1", False, False),
    ("Wang Xuan", "2", False, False),
    ("Zhang Shufang", "1", False, True),
    ("Xiang Yang", "2", False, True),
]

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for i, (name, aff_num, co_first, co_corr) in enumerate(authors_data):
    if i > 0:
        add_run(p, ", ", 12)
    add_run(p, name, 12)
    add_superscript_run(p, aff_num, 9)
    if co_first:
        add_superscript_run(p, "†", 9)  # dagger
    if co_corr:
        add_superscript_run(p, "*", 9)

# ── Affiliations ──
affiliations = {
    "1": "Central Laboratory, Haikou Affiliated Hospital of Central South University Xiangya School of Medicine, Haikou 570208, China",
    "2": "Department of Hepatobiliary Surgery, Haikou Affiliated Hospital of Central South University Xiangya School of Medicine, Haikou 570208, China",
    "3": "Department of Anesthesiology, Changde Hospital, Xiangya School of Medicine, Central South University, Changde 415000, China",
}
for num in ["1", "2", "3"]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_superscript_run(p, num, 9)
    add_run(p, " " + affiliations[num], 9, italic=True)

# Co-first author note
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_superscript_run(p, "†", 9)
add_run(p, " Lei Zhongwen and Wang Jinnong contributed equally to this work.", 9, italic=True)

# Corresponding author
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p, "*Correspondence: ", 9)
add_run(p, "Xiang Yang, Email: xiangyang200611@126.com; Zhang Shufang, Email: zsf66189665@126.com", 9)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_run(p2, "Haikou Affiliated Hospital of Central South University Xiangya School of Medicine, Haikou 570208, China", 9, italic=True)

doc.add_paragraph()

# ── Abstract ──
add_heading(doc, "Abstract", level=1)
abstract = (
    "Cholangiocarcinoma (CCA) is an aggressive biliary malignancy with limited therapeutic options. "
    "Cancer-associated fibroblasts (CAFs) and tumor-associated macrophages (TAMs) contribute to CCA progression, "
    "but systematic multi-omics characterization of CAF–TAM crosstalk remains limited. "
    "We integrated bulk transcriptomic data from TCGA-CHOL (n=44) and GSE107943 (n=57), "
    "single-cell RNA-seq from GSE138709 (32,626 cells, 5 intrahepatic CCA patients), "
    "and expression validation in GSE26566 (n=169). "
    "Cross-cohort analysis identified 344 immunometabolic, CAF, and TAM-related differentially expressed genes "
    "validated with 99.4% directional concordance. "
    "Upregulated genes were enriched in extracellular matrix organization and collagen metabolism, "
    "while downregulated genes were enriched in lipid, steroid, and amino acid metabolic processes. "
    "An aggressive microenvironment score (ECM_up + CAF + TAM − Metabolism_down) correlated positively "
    "with CAF abundance (rho=0.77), M2 macrophage markers (rho=0.65), and immune checkpoint transcript levels "
    "including HAVCR2 (rho=0.54), IDO1 (rho=0.48), and CD86 (rho=0.55) in TCGA tumors. "
    "Single-cell analysis localized the highest aggressive score to CAF_Fibroblast cells and identified "
    "Macrophage_TAM as the predominant transcript-level expresser of HAVCR2, IDO1, CD86, and PDCD1LG2. "
    "CellChat analysis inferred communication axes involving COL1A1/COL1A2–CD44 and MIF–CD74/CXCR4 "
    "between CAFs, TAMs, and epithelial cells. "
    "Molecular docking provided in silico support for target–ligand feasibility "
    "(CXCR4–Plerixafor: −8.01 kcal/mol; IDO1–Epacadostat: −6.47 kcal/mol). "
    "These findings characterize a CAF–TAM communication-associated aggressive microenvironment in CCA "
    "and nominate candidate targets warranting further experimental investigation."
)
add_para(doc, abstract, font_size=11)

# Keywords
p = doc.add_paragraph()
add_run(p, "Keywords: ", 11, bold=True)
add_run(p, "cholangiocarcinoma; tumor microenvironment; cancer-associated fibroblasts; "
        "tumor-associated macrophages; single-cell RNA-seq; CellChat; immunometabolism; "
        "molecular docking", 11)

# ═══════════════════════════════════════════════════════════════
# BODY SECTIONS
# ═══════════════════════════════════════════════════════════════

with open(MANUSCRIPT / "draft_v8_CSBJ_formatted.md", "r", encoding="utf-8") as f:
    v8 = f.read()

body_start = v8.find("## Introduction")
lines = v8[body_start:].split("\n")

current_section = None
current_text = []
section_texts = {}

for line in lines:
    m = re.match(r'^## (.+)$', line)
    if m:
        if current_section:
            section_texts[current_section] = "\n".join(current_text)
        current_section = m.group(1).strip()
        current_text = []
        continue
    m2 = re.match(r'^### (.+)$', line)
    if m2:
        current_text.append(f"SUBHEADING:{m2.group(1).strip()}")
        continue
    current_text.append(line)
if current_section:
    section_texts[current_section] = "\n".join(current_text)

ordered = [
    "Introduction", "Materials and Methods", "Results", "Discussion",
    "Data Availability", "Code Availability", "Ethics Statement",
    "Author Contributions", "Funding", "Conflict of Interest", "Acknowledgments",
    "References", "Figure Legends", "Supplementary Materials",
]

AUTHOR_CONTRIBUTIONS = [
    ("Lei Zhongwen", "Conceptualization, Methodology, Software, Formal Analysis, Investigation, Data Curation, Visualization, Writing – Original Draft."),
    ("Wang Jinnong", "Methodology, Formal Analysis, Validation, Data Curation, Visualization, Writing – Original Draft."),
    ("Gao Yuanhui", "Investigation, Data Curation, Validation, Writing – Review & Editing."),
    ("Huang Denggao", "Data Curation, Visualization, Validation, Writing – Review & Editing."),
    ("Wang Xuan", "Resources, Investigation, Writing – Review & Editing."),
    ("Zhang Shufang", "Conceptualization, Resources, Supervision, Project Administration, Writing – Review & Editing."),
    ("Xiang Yang", "Conceptualization, Resources, Supervision, Project Administration, Writing – Review & Editing."),
]

# Figure legends (embedded directly, not parsed from markdown)
FIGURE_LEGENDS = [
    ("Figure 1. Study design and analytical workflow.",
     "Schematic overview of the integrated computational framework combining bulk transcriptomic cohorts (TCGA-CHOL and GSE107943), single-cell RNA-seq (GSE138709), external expression validation (GSE26566), cross-cohort differential expression analysis, ssGSEA-based microenvironment scoring, immune and checkpoint correlation analysis, CellChat-inferred communication analysis, hub gene prioritization, druggability assessment, and in silico molecular docking. CellChat-based communication axes and docking scores represent computational predictions requiring experimental validation."),
    ("Figure 2. Cross-cohort identification and functional characterization of validated IM/CAF/TAM-related DEGs.",
     "(A–D) PCA and volcano plots for TCGA-CHOL and GSE107943 tumor–normal comparisons. (E) Overlap of significant DEGs between TCGA-CHOL and GSE107943 with concordant direction of dysregulation. (F) GO BP enrichment of upregulated validated IM/CAF/TAM DEGs. (G) GO BP enrichment of downregulated validated IM/CAF/TAM DEGs. (H) Summary of upregulated and downregulated pathway enrichment patterns."),
    ("Figure 3. Aggressive microenvironment score and immune/checkpoint associations.",
     "(A–B) ssGSEA pathway score patterns in TCGA-CHOL and GSE107943. (C) Exploratory survival analysis of the aggressive microenvironment score in GSE107943. All survival analyses are exploratory. (D) Correlation structure among microenvironment-related scores. (E) Correlation between aggressive score and immune cell marker-based ssGSEA scores. (F) Correlation between aggressive score and checkpoint transcript expression. (G–H) Representative scatter plots showing associations between aggressive score and macrophage/checkpoint-related features."),
    ("Figure 4. Single-cell localization of the aggressive microenvironment program.",
     "(A) UMAP visualization of GSE138709 cells by unsupervised clusters. (B) UMAP visualization after marker-based cell type annotation. (C) Single-cell aggressive microenvironment score projected onto UMAP. (D) Distribution of aggressive score across annotated cell types. (E) Dot plot of selected CAF, TAM, checkpoint, and metabolic marker genes. Checkpoint expression patterns are transcript-level observations and require protein-level validation."),
    ("Figure 5. CellChat-inferred CAF–TAM–Epithelial communication network.",
     "(A) Simplified directed network summarizing inferred communication among CAF_Fibroblast, Macrophage_TAM, and Epithelial_like cells. Edge labels indicate ligand–receptor pair counts and representative axes. (B) Incoming and outgoing signaling patterns across tumor cell types. (C) Bubble plot of inferred CAF-to-TAM ligand–receptor pairs. (D) Bubble plot of inferred TAM-to-CAF ligand–receptor pairs. CellChat results represent computational inference based on ligand–receptor co-expression and do not establish functional signaling."),
    ("Figure 6. Integrated hub gene prioritization and exploratory clinical relevance.",
     "(A) Multi-layer evidence heatmap for candidate hub genes. (B) Ranking of top hub genes based on integrated evidence scores. (C) Prioritized ligand–receptor communication axes. (D) Exploratory association between CAF–TAM axis scores and overall survival in GSE107943. These survival analyses are exploratory and require validation in larger independent cohorts."),
    ("Figure 7. Druggability assessment and in silico docking of candidate targets.",
     "(A) Druggability classification of prioritized targets. (B) Predicted docking affinities for selected target–ligand pairs. Docking scores are computational predictions and do not represent experimental binding affinities. (C) Proposed model summarizing a CAF–TAM communication-associated aggressive microenvironment in cholangiocarcinoma and candidate targets for future validation."),
]

for sec_name in ordered:
    if sec_name == "Author Contributions":
        add_heading(doc, sec_name, level=1)
        for name, contrib in AUTHOR_CONTRIBUTIONS:
            p = doc.add_paragraph()
            add_run(p, f"{name}: ", 11, bold=True)
            add_run(p, contrib, 11)
        doc.add_paragraph()
        add_para(doc, "Lei Zhongwen and Wang Jinnong contributed equally to this work.", italic=True, font_size=10)
        add_para(doc, "All authors read and approved the final manuscript.", font_size=10)

    elif sec_name == "Funding":
        add_heading(doc, sec_name, level=1)
        add_para(doc, "This research did not receive any specific grant from funding agencies in the public, commercial, or not-for-profit sectors.", font_size=11)

    elif sec_name == "Conflict of Interest":
        add_heading(doc, sec_name, level=1)
        add_para(doc, "The authors declare that they have no known competing financial interests or personal relationships that could have appeared to influence the work reported in this paper.", font_size=11)

    elif sec_name == "Acknowledgments":
        add_heading(doc, sec_name, level=1)
        add_para(doc, "The authors thank all investigators who generated and shared the public datasets used in this study.", font_size=11)

    elif sec_name == "Data Availability":
        add_heading(doc, sec_name, level=1)
        add_para(doc, "All raw datasets analyzed in this study are publicly available from GDC "
                 "(TCGA-CHOL) and GEO (GSE107943, GSE138709, GSE26566). Processed intermediate "
                 "tables generated during this study are available from the corresponding authors "
                 "upon reasonable request.", font_size=11)

    elif sec_name == "Code Availability":
        add_heading(doc, sec_name, level=1)
        add_para(doc, "The analysis scripts used in this study (18 R + 2 Python) are documented "
                 "with headers specifying purpose, input, output, and methods. All analysis scripts "
                 "are available from the corresponding authors upon reasonable request and will be "
                 "deposited in a public repository upon acceptance.", font_size=11)

    elif sec_name == "Ethics Statement":
        add_heading(doc, sec_name, level=1)
        add_para(doc, "This study exclusively used publicly available, de-identified data from GDC "
                 "and GEO. No new human subjects research was conducted. All original studies obtained "
                 "appropriate ethical approvals as described in their respective publications.", font_size=11)

    elif sec_name == "References":
        add_heading(doc, sec_name, level=1)
        text = section_texts.get(sec_name, "")
        for line in text.split("\n"):
            m = re.match(r'^\[(\d+)\]\s+(.+)$', line)
            if m:
                add_ref(doc, int(m.group(1)), m.group(2).strip())

    elif sec_name == "Figure Legends":
        add_heading(doc, sec_name, level=1)
        for fig_title, fig_body in FIGURE_LEGENDS:
            p = doc.add_paragraph()
            add_run(p, fig_title, 11, bold=True)
            add_run(p, " " + fig_body, 11)
        print("  Figure Legends: 7 figures embedded")

    elif sec_name == "Supplementary Materials":
        add_heading(doc, sec_name, level=1)
        text = section_texts.get(sec_name, "")
        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("SUBHEADING:"):
                add_heading(doc, line.replace("SUBHEADING:", "").strip(), level=2)
            elif line.startswith("- "):
                add_para(doc, line, font_size=10)
            elif line and not line.startswith("---"):
                if line.startswith("**"):
                    p = doc.add_paragraph()
                    add_run(p, line.replace("**", ""), 11, bold=True)
                else:
                    add_para(doc, line, font_size=11)

    elif sec_name in section_texts:
        add_heading(doc, sec_name, level=1)
        text = section_texts[sec_name]

        # FIX: Replace Supplementary Table S10 → S17 in Results
        if sec_name == "Results":
            text = text.replace(
                "(Supplementary Table S10)",
                "(Supplementary Table S17)")
            print("  S10 -> S17 fix applied in Results")

        for line in text.split("\n"):
            line = line.strip()
            if line.startswith("SUBHEADING:"):
                add_heading(doc, line.replace("SUBHEADING:", "").strip(), level=2)
            elif line and not line.startswith("---"):
                if line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    add_run(p, line.replace("**", ""), 11, bold=True)
                elif line.strip():
                    add_para(doc, line, font_size=11)

# Save
path = DST / "Manuscript_CSBJ_final_clean.docx"
doc.save(str(path))
print(f"Saved: {path.name}")
