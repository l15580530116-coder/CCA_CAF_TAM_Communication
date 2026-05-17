#!/usr/bin/env python
"""Step 40: Redraw Figure 7B as dotmap, fix manuscript text issues, generate v3."""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.gridspec as gridspec
from matplotlib.patches import FancyBboxPatch, Rectangle
import numpy as np
from pathlib import Path
import fitz, csv, io, shutil
from PIL import Image

BASE = Path(r"e:\CCA")
SRC = BASE / "figures"
TBL = BASE / "tables"
SRC_FIG7 = BASE / "submission_CSBJ_UPLOAD_READY_FINAL/Figures_Main_FINAL"
OUT = BASE / "figures/main_final_v7"
OUT.mkdir(parents=True, exist_ok=True)
DPI = 600
LF = 18; NF = 8

def pdf_img(path, pg=0, dpi=DPI, max_mp=22):
    doc = fitz.open(str(path))
    if len(doc)==0: doc.close(); raise ValueError("0-page")
    if pg>=len(doc): pg=len(doc)-1
    page = doc[pg]; pw,ph = page.rect.width, page.rect.height
    d=dpi
    if (pw*d/72)*(ph*d/72)/1e6>max_mp: d=int(dpi*(max_mp/((pw*dpi/72)*(ph*dpi/72)/1e6))**0.5)
    zoom=d/72.0; mat=fitz.Matrix(zoom,zoom)
    pix=page.get_pixmap(matrix=mat,colorspace=fitz.csRGB)
    img=np.array(Image.open(io.BytesIO(pix.tobytes("png")))); doc.close()
    return img

def add_label(ax, letter, fs=LF):
    ax.text(0.012,0.978,letter,transform=ax.transAxes,fontsize=fs,fontweight="bold",
            color="black",va="top",ha="left",
            bbox=dict(boxstyle="round,pad=0.08",facecolor="white",edgecolor="none",alpha=0.88))

def add_note(ax, text, fs=NF):
    ax.text(0.985,0.015,text,transform=ax.transAxes,fontsize=fs,fontstyle="italic",
            color="#666666",va="bottom",ha="right",
            bbox=dict(boxstyle="round,pad=0.08",facecolor="white",edgecolor="none",alpha=0.88))

# ═══════════════════════════════════════════════════════════════
# 1. REDRAW FIGURE 7B AS DRUG-TARGET DOTMAP
# ═══════════════════════════════════════════════════════════════
print("=== Redrawing Figure 7B as drug-target dotmap ===")

with open(TBL/"drug_screening/candidate_drugs_prioritized.csv",newline="") as f:
    drugs = list(csv.DictReader(f))

# Get unique targets and drugs (top 10 targets, top 10 drugs)
tgt_counts = {}
for r in drugs:
    t = r["target_gene"]; tgt_counts[t] = tgt_counts.get(t,0)+1
top_targets = [t for t,c in sorted(tgt_counts.items(),key=lambda x:-x[1])][:10]

# Get unique drugs in order
drugs_seen = []
for r in drugs:
    d = r["drug_name"]
    if r["target_gene"] in top_targets and d not in drugs_seen:
        drugs_seen.append(d)
drugs_seen = drugs_seen[:10]

# Build matrix
matrix = np.zeros((len(top_targets), len(drugs_seen)))
for r in drugs:
    t = r["target_gene"]; d = r["drug_name"]
    if t in top_targets and d in drugs_seen:
        ti = top_targets.index(t); di = drugs_seen.index(d)
        matrix[ti, di] = 1

fig_b, ax_b = plt.subplots(1,1,figsize=(10,6.5),facecolor="white")
# Draw as scatter/dotmap
rows, cols = np.where(matrix > 0)
ax_b.scatter(cols+0.5, rows+0.5, s=180, c="#E74C3C", marker="s", edgecolors="white", linewidth=0.5, zorder=3, alpha=0.85)

ax_b.set_xticks(np.arange(len(drugs_seen))+0.5)
ax_b.set_xticklabels(drugs_seen, fontsize=7, rotation=45, ha="right")
ax_b.set_yticks(np.arange(len(top_targets))+0.5)
ax_b.set_yticklabels(top_targets, fontsize=8.5, fontweight="bold")
ax_b.set_xlim(0, len(drugs_seen)); ax_b.set_ylim(0, len(top_targets))
ax_b.invert_yaxis()
# Grid lines
for i in range(len(top_targets)+1): ax_b.axhline(i, color="#DDDDDD", lw=0.5, zorder=0)
for j in range(len(drugs_seen)+1): ax_b.axvline(j, color="#DDDDDD", lw=0.5, zorder=0)
ax_b.set_title("Candidate drug-target relationships", fontsize=12, fontweight="bold")
plt.tight_layout()

import io as io_mod
fig_b.canvas.draw()
img_data = np.frombuffer(fig_b.canvas.buffer_rgba(), dtype=np.uint8)
w, h = fig_b.canvas.get_width_height()
img_b = img_data.reshape(h, w, 4)[:,:,:3]
plt.close(fig_b)
print(f"  Panel B: dotmap ({img_b.shape[1]}x{img_b.shape[0]})")

# ═══════════════════════════════════════════════════════════════
# 2. RE-ASSEMBLE FIGURE 7 WITH NEW PANEL B
# ═══════════════════════════════════════════════════════════════
print("\n=== Re-assembling Figure 7 ===")

# Reuse panels A, C, D, E, F from v6 logic
# Panel A: existing druggability
img_a = pdf_img(SRC/"drug_screening/Fig11A_target_druggability_heatmap.pdf")

# Panel C: drugs per target (recreate from data)
tgt_counts_sorted = sorted(tgt_counts.items(), key=lambda x:-x[1])[:12]
fig_c, ax_c = plt.subplots(1,1,figsize=(7,6),facecolor="white")
tl=[t for t,c in tgt_counts_sorted]; cl=[c for t,c in tgt_counts_sorted]
colors_c=["#E74C3C" if c>=3 else "#3498DB" if c>=2 else "#95A5A6" for c in cl]
ax_c.barh(range(len(tl)),cl,color=colors_c,edgecolor="white",height=0.7)
ax_c.set_yticks(range(len(tl))); ax_c.set_yticklabels(tl,fontsize=8)
ax_c.set_xlabel("Candidate drugs",fontsize=10)
ax_c.set_title("Drugs per target",fontsize=11,fontweight="bold"); ax_c.invert_yaxis()
for i,(t,c) in enumerate(zip(tl,cl)): ax_c.text(c+0.05,i,str(c),va="center",fontsize=9,fontweight="bold")
ax_c.set_xlim(0,max(cl)+1.5); plt.tight_layout()
fig_c.canvas.draw(); wc,hc=fig_c.canvas.get_width_height()
img_c=np.frombuffer(fig_c.canvas.buffer_rgba(),dtype=np.uint8).reshape(hc,wc,4)[:,:,:3]; plt.close(fig_c)

# Panel D: docking affinity
img_d = pdf_img(SRC/"docking/FigS_Docking_affinity_barplot_all_pairs.pdf")

# Panel E: residues
with open(TBL/"docking/docking_neighbor_residues_4A.csv",newline="") as f:
    residues=list(csv.DictReader(f))
pc={}; [pc.update({r["pair"]:pc.get(r["pair"],0)+1}) for r in residues]
sp=sorted(pc.items(),key=lambda x:-x[1])
fig_e,ax_e=plt.subplots(1,1,figsize=(7,5.5),facecolor="white")
pn=[p.replace("_","-") for p,c in sp]; rc=[c for p,c in sp]
ax_e.barh(range(len(pn)),rc,color=["#9B59B6","#8E44AD","#7D3C98","#6C3483"],edgecolor="white",height=0.65)
ax_e.set_yticks(range(len(pn))); ax_e.set_yticklabels(pn,fontsize=8)
ax_e.set_xlabel("Residues within 4 A",fontsize=10)
ax_e.set_title("Docking neighbor residues",fontsize=11,fontweight="bold"); ax_e.invert_yaxis()
for i,c in enumerate(rc): ax_e.text(c+0.3,i,str(c),va="center",fontsize=9,fontweight="bold")
ax_e.set_xlim(0,max(rc)+3); plt.tight_layout()
fig_e.canvas.draw(); we,he=fig_e.canvas.get_width_height()
img_e=np.frombuffer(fig_e.canvas.buffer_rgba(),dtype=np.uint8).reshape(he,we,4)[:,:,:3]; plt.close(fig_e)

# Panel F: large mechanism
fig_f,ax_f=plt.subplots(1,1,figsize=(20,9),facecolor="white")
ax_f.set_xlim(0,20); ax_f.set_ylim(0,11); ax_f.axis("off")
cells=[("CAF\nFibroblast",3.5,7.0,"#E74C3C"),("Macrophage\nTAM",10.0,7.0,"#3498DB"),("Epithelial\nlike",16.5,7.0,"#27AE60")]
for label,x,y,color in cells:
    box=FancyBboxPatch((x-2.0,y-1.0),4.0,2.0,boxstyle="round,pad=0.1,rounding_size=0.2",facecolor=color,edgecolor="white",lw=2.5,alpha=0.22)
    ax_f.add_patch(box); ax_f.text(x,y,label,ha="center",va="center",fontsize=12,fontweight="bold",color=color)
arrows=[(5.5,7.2,8.0,7.2,"COL1A1/COL1A2 -> CD44"),(5.5,6.5,8.0,6.5,"MIF -> CD74/CXCR4"),
        (8.0,5.8,5.5,5.8,"TGFB1 -> TGFBR"),(8.0,5.3,5.5,5.3,"PPIA -> BSG"),
        (16.5,8.0,12.0,8.0,"MIF -> CD74/CXCR4")]
for x1,y1,x2,y2,label in arrows:
    ax_f.annotate("",xy=(x2,y2),xytext=(x1,y1),arrowprops=dict(arrowstyle="->",color="#555555",lw=2.5))
    mx,my=(x1+x2)/2,(y1+y2)/2+0.2
    ax_f.text(mx,my,label,ha="center",va="bottom",fontsize=8,color="#333333",
              bbox=dict(boxstyle="round,pad=0.06",facecolor="white",edgecolor="none",alpha=0.85))
tam_box=FancyBboxPatch((7.0,3.5),6.0,1.3,boxstyle="round,pad=0.1,rounding_size=0.12",facecolor="#3498DB",edgecolor="white",lw=1.5,alpha=0.15)
ax_f.add_patch(tam_box); ax_f.text(10.0,4.3,"TAM checkpoint transcripts (mRNA-level)",ha="center",fontsize=9.5,fontweight="bold",color="#2980B9")
ax_f.text(10.0,3.8,"HAVCR2  |  IDO1  |  CD86  |  PDCD1LG2",ha="center",fontsize=8.5,color="#333333")
tgt_box=FancyBboxPatch((7.0,2.0),6.0,1.0,boxstyle="round,pad=0.08,rounding_size=0.10",facecolor="#27AE60",edgecolor="white",lw=1.5,alpha=0.15)
ax_f.add_patch(tgt_box); ax_f.text(10.0,2.6,"Candidate targets: CXCR4  |  IDO1  |  TGFBR1  |  MIF",ha="center",fontsize=9,fontweight="bold",color="#1E8449")
ax_f.text(10.0,2.25,"Docking = in silico predictions only",ha="center",fontsize=7,fontstyle="italic",color="#888888")
import matplotlib.patches as mpatches
ax_f.legend(handles=[mpatches.Patch(color="#E74C3C",alpha=0.22,label="CAF Fibroblast"),
                      mpatches.Patch(color="#3498DB",alpha=0.22,label="Macrophage TAM"),
                      mpatches.Patch(color="#27AE60",alpha=0.22,label="Epithelial-like")],
            loc="upper right",fontsize=7.5,framealpha=0.9)
ax_f.text(10.0,10.3,"Proposed CAF-TAM-Epithelial communication model",ha="center",fontsize=14,fontweight="bold")
ax_f.text(10.0,0.6,"Computational model; experimental validation required.",ha="center",fontsize=8,fontstyle="italic",color="#888888")
plt.tight_layout()
fig_f.canvas.draw(); wf,hf=fig_f.canvas.get_width_height()
img_f=np.frombuffer(fig_f.canvas.buffer_rgba(),dtype=np.uint8).reshape(hf,wf,4)[:,:,:3]; plt.close(fig_f)

# Assemble
fig7=plt.figure(figsize=(20,19),facecolor="white")
gs=gridspec.GridSpec(3,6,figure=fig7,wspace=0.08,hspace=0.12,
                     left=0.025,right=0.975,top=0.985,bottom=0.015,
                     height_ratios=[1.0,1.0,1.6])
ax_a=fig7.add_subplot(gs[0,0:3]); ax_a.imshow(img_a); ax_a.axis("off"); add_label(ax_a,"A")
ax_b2=fig7.add_subplot(gs[0,3:6]); ax_b2.imshow(img_b); ax_b2.axis("off"); add_label(ax_b2,"B")
ax_c2=fig7.add_subplot(gs[1,0:2]); ax_c2.imshow(img_c); ax_c2.axis("off"); add_label(ax_c2,"C")
ax_d2=fig7.add_subplot(gs[1,2:4]); ax_d2.imshow(img_d); ax_d2.axis("off"); add_label(ax_d2,"D"); add_note(ax_d2,"In silico predictions only")
ax_e2=fig7.add_subplot(gs[1,4:6]); ax_e2.imshow(img_e); ax_e2.axis("off"); add_label(ax_e2,"E")
ax_f2=fig7.add_subplot(gs[2,:]); ax_f2.imshow(img_f); ax_f2.axis("off"); add_label(ax_f2,"F",fs=LF); add_note(ax_f2,"Computational model; experimental validation required")
fig7.suptitle("Figure 7. Candidate target prioritization and in silico docking",fontsize=14,fontweight="bold",y=0.998)
for ext in ["pdf","png"]:
    p=OUT/f"Figure7_therapeutic_implications_v7.{ext}"; fig7.savefig(p,dpi=DPI,facecolor="white",edgecolor="none",bbox_inches="tight")
    print(f"  {p.name} ({p.stat().st_size//1024} KB)")
plt.close(fig7)

# Copy to upload destination
for ext in ["pdf","png"]:
    shutil.copy2(OUT/f"Figure7_therapeutic_implications_v7.{ext}",
                 SRC_FIG7/f"Figure7_therapeutic_implications.{ext}")
print("  Copied to Figures_Main_FINAL/")

# ═══════════════════════════════════════════════════════════════
# 3. FIX MANUSCRIPT TEXT
# ═══════════════════════════════════════════════════════════════
print("\n=== Fixing manuscript text ===")
from docx import Document
from docx.shared import Pt

doc = Document(str(BASE/"submission_CSBJ_UPLOAD_READY_FINAL/Manuscript_CSBJ_final_clean_v2.docx"))

# Collect all paragraph text first
paras = [(i, p.text) for i, p in enumerate(doc.paragraphs)]

fixes_applied = 0

for idx, para in enumerate(doc.paragraphs):
    text = para.text

    # Fix 1: Figure 5 legend - remove duplicate "The top inferred..." sentence
    if text.startswith("Figure 5.") and "The top inferred signaling pathways and source-target communication strengths are summarized in Figure 5E and Figure 5F." in text:
        # This sentence should NOT be in the legend - remove it
        new_text = text.replace(
            ". The top inferred signaling pathways and source-target communication strengths are summarized in Figure 5E and Figure 5F.",
            ".")
        # Also handle case where it appears without leading period
        new_text = new_text.replace(
            "The top inferred signaling pathways and source-target communication strengths are summarized in Figure 5E and Figure 5F.",
            "")
        new_text = new_text.replace("..", ".")
        new_text = new_text.strip()
        if para.runs:
            para.runs[0].text = new_text
            for r in para.runs[1:]: r.text = ""
        fixes_applied += 1
        print("  FIXED: Figure 5 legend (removed duplicate sentence)")

    # Fix 2: Results 4.7 double period
    if "precluding survival validation in this third cohort.." in text:
        new_text = text.replace(
            "precluding survival validation in this third cohort..",
            "precluding survival validation in this third cohort.")
        if para.runs:
            # Find and fix in the last run
            for run in para.runs:
                if "precluding survival validation in this third cohort.." in run.text:
                    run.text = run.text.replace("..", ".")
                    break
        fixes_applied += 1
        print("  FIXED: double period in Results 4.7")

    # Fix 3: Results 4.8 - replace docking paragraph
    if "Molecular docking was performed using AutoDock Vina v1.2.7" in text and "(Figure 7B)" in text:
        # Find runs containing the wrong reference
        for run in para.runs:
            if "(Figure 7B)" in run.text:
                run.text = run.text.replace("(Figure 7B)", "(Figure 7D)")
                fixes_applied += 1
                print("  FIXED: Figure 7B -> Figure 7D for grid centers")

        # Fix double period after Supplementary Table S17
        for run in para.runs:
            if "Supplementary Table S17).." in run.text:
                run.text = run.text.replace("Supplementary Table S17)..",
                                            "Supplementary Table S17).")
                fixes_applied += 1
                print("  FIXED: double period after S17 reference")

    # Fix 4: Remove duplicate 43 residues sentence
    # This is complex to fix programmatically - will handle via the docking paragraph replacement below

# Now handle the Results 4.8 docking paragraph comprehensively
# Find the paragraph that contains the docking results
for idx, para in enumerate(doc.paragraphs):
    text = para.text
    if "Molecular docking was performed using AutoDock Vina v1.2.7" in text and "43 residues within 4" in text:
        # This is the docking paragraph. Replace it with the clean version.
        clean_docking = (
            "Molecular docking was performed using AutoDock Vina v1.2.7 [23,24] "
            "for four target-ligand pairs with available co-crystal structures. "
            "Ligand preparation was performed with Open Babel v3.1.1 [25]. "
            "Candidate drug-target relationships are summarized in Figure 7B, "
            "with candidate drug counts per target shown in Figure 7C. "
            "The best predicted binding affinities were: "
            "CXCR4-Plerixafor (-8.01 kcal/mol, PDB: 3ODU), "
            "IDO1-Epacadostat (-6.47 kcal/mol, PDB: 5WN8), "
            "MIF-ISO-1 (-5.08 kcal/mol, PDB: 1LJT), and "
            "TGFBR1-Galunisertib (-4.83 kcal/mol, PDB: 6B8Y) (Figure 7D). "
            "Grid centers were defined by co-crystal ligand centroids. "
            "Docking neighbor residue analysis identified 43 residues within 4 A "
            "of the best docking poses across the four target-ligand pairs "
            "(Figure 7E; Supplementary Table S17). "
            "A proposed CAF-TAM-Epithelial model integrating these findings "
            "is shown in Figure 7F."
        )
        # Only replace if different from current
        if para.runs:
            para.runs[0].text = clean_docking
            for r in para.runs[1:]: r.text = ""
            fixes_applied += 1
            print("  FIXED: Results 4.8 docking paragraph (removed duplicate, fixed refs)")

    # Fix 5: Update Figure 7 legend
    if text.startswith("Figure 7.") and "Candidate drug-target relationship network" in text:
        new_legend = (
            "Figure 7. Candidate target prioritization and in silico docking. "
            "(A) Druggability classification of prioritized targets. "
            "(B) Candidate drug-target relationship matrix. "
            "(C) Candidate drug count per target. "
            "(D) Predicted docking affinities for selected target-ligand pairs. "
            "(E) Docking neighbor residue summary. "
            "(F) Proposed model summarizing a CAF-TAM communication-associated aggressive "
            "microenvironment in cholangiocarcinoma and candidate targets for future validation. "
            "Docking scores are computational predictions and do not represent experimental binding affinities."
        )
        if para.runs:
            para.runs[0].text = new_legend
            for r in para.runs[1:]: r.text = ""
            fixes_applied += 1
            print("  FIXED: Figure 7 legend (network -> matrix)")

print(f"\nTotal fixes applied: {fixes_applied}")

# Save v3
dst = BASE/"submission_CSBJ_UPLOAD_READY_FINAL/Manuscript_CSBJ_final_clean_v3.docx"
doc.save(str(dst))
print(f"Saved: {dst.name}")
